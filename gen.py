import auxil
import consts

from loguru import logger
from random import choice, randint
from docx import Document
import json
import argparse
from os import makedirs
import re

from sys import exit #delete from prod

def load_samples(samples_dir):
	with open(samples_dir + "/headers.txt") as headerfile:
		header = headerfile.read().split(";;\n")

	with open(samples_dir + "/names.txt") as namefile:
		name = namefile.read().split(";;\n")

	with open(samples_dir + "/intros.txt") as introfile:
		intro = introfile.read().split(";;\n")

	with open(samples_dir + "/instructions.txt") as instructionfile:
		instruction = instructionfile.read().split(";;\n")

	with open(samples_dir + "/responsible.txt") as responsiblefile:
		responsible = responsiblefile.read().split(";;\n")

	with open(samples_dir + "/creators.txt") as creatorfile:
		creator = creatorfile.read().split('\n')

	logger.info(f"[header] length: {len(header)}")
	logger.info(f"[name] length: {len(name)}")
	logger.info(f"[intro] length: {len(intro)}")
	logger.info(f"[instruction] length: {len(instruction)}")
	logger.info(f"[responsible] length: {len(responsible)}")
	logger.info(f"[creator] length: {len(creator)}")
	vars = len(header)*len(name)*len(intro)*len(instruction)
	logger.info(f"Maximum number of decrees: {vars*len(responsible)*len(creator)}")
	logger.warning(f"Approximate number of different decrees: {vars}")

	return (header, name, intro, instruction, responsible, creator)

def add_numbering(instruction):
	clauses = re.split(r"\{\d*\}", instruction)[1:]
	numbering_types = [choice(consts.numbering_types) for _ in range(3)]
	complete_instruction = [{"clause": clauses[0],
							"index": 1,
							"nesting_level": 0,
							"numbering_type": numbering_types[0]}]
	numbering = [1, 1, 1]

	for indx in range(1, len(clauses)):
		clause = clauses[indx]
		prev_clause = complete_instruction[indx-1]
		prev_nesting_level = prev_clause["nesting_level"]
		nesting_level = randint(0, 1) if prev_nesting_level == 0 else randint(0, 2)
		n_type = numbering_types[nesting_level]

		if prev_nesting_level == nesting_level:
			numbering[nesting_level] += 1
			index = numbering[nesting_level]
		elif prev_nesting_level < nesting_level:
			numbering[nesting_level] = 1
			index = numbering[nesting_level]
		elif prev_nesting_level > nesting_level:
			numbering[nesting_level] += 1
			index = numbering[nesting_level]

		complete_instruction.append({"clause": clause,
									 "index": index,
									 "nesting_level": nesting_level,
									 "numbering_type": n_type})

	for indx in range(len(clauses)):
		index = complete_instruction[indx]["index"]
		nesting_level = complete_instruction[indx]["nesting_level"]
		n_type = complete_instruction[indx]["numbering_type"]
		if n_type[0] == "arabic":
			clauses[indx] = '\t'*nesting_level + str(index) + n_type[1] + clauses[indx]

		elif n_type[0] == "roman":
			clauses[indx] = '\t'*nesting_level + str(auxil.to_roman(index)) + n_type[1] + clauses[indx]

		elif n_type[0] == "bullet":
			clauses[indx] = '\t'*nesting_level + n_type[1] + clauses[indx]

		elif n_type[0] == "latin":
			clauses[indx] = '\t'*nesting_level + consts.latin_alphabet[index-1] + n_type[1] + clauses[indx]

	instruction = "".join(clauses)

	return instruction

def write_docx(header, name, intro, instruction,
	responsible, creator, date, out, count):
	document = Document()

	headerp = document.add_paragraph()
	headerp.alignment = 1
	headerp.add_run(header + '\n\n').bold = True
	
	namep = document.add_paragraph()
	namep.alignment = 1
	namep.add_run(name)

	document.add_paragraph(intro)

	
	document.add_paragraph(instruction)

	responsiblep = document.add_paragraph("Ответственным за исполнение настоящего приказа назначить ")
	responsiblep.add_run(responsible)

	datep = document.add_paragraph(creator+'\n')
	datep.add_run(date)
	datep.alignment = 2

	document.save(f"{out}/docx/{count}.docx")
	logger.info(f"Saved {out}/docx/{count}.docx")

def write_json(instruction, responsible, date, out, count):
	json_dict = {"instruction": instruction, "responsible": responsible, "date": date}

	with open(f"{out}/json/{count}.json", "w") as jsonf:
		json.dump(json_dict, jsonf, ensure_ascii=False, indent=4)
		logger.info(f"Saved {out}/json/{count}.json")


def generate(data, out, size):
	try:
		makedirs(out+"/docx")
		makedirs(out+"/json")
		logger.info(f"{out}/docx and {out}/json was created.")
	except FileExistsError:
		pass

	count = 0
	while True:
		header = choice(data[0])
		name = choice(data[1])
		intro = choice(data[2])
		instruction = add_numbering(choice(data[3]))
		responsible = choice(data[4])
		creator = choice(data[5])
		date = auxil.generate_date()

		write_docx(header, name, intro, instruction,
			responsible, creator, date, out, count)

		write_json(instruction, responsible, date, out, count)

		count += 1
		if count % 25 == 0:
			if auxil.getsize(out) >= size:
				logger.warning(f"Size of {out} dir: {auxil.getsize(out)} bytes")
				break

def get_args():
	parser = argparse.ArgumentParser()
	parser.add_argument("size", help="Max size of output dir. For example: 10KB, 10MB, 10GB",
						type=auxil.check_size_format)
	parser.add_argument("-s", "--samples", help="Path to dir with samples",
						metavar="path", type=str, default="./samples/")
	parser.add_argument("-o", "--out", help="Path for output files",
						metavar="path", type=str, default="./")
	parser.add_argument("-v", "--verbose", action="store_true")
	return parser.parse_args()

def main():
	args = get_args()
	auxil.logger_config(args.verbose)

	data = load_samples(args.samples)
	bytes_size = auxil.size_to_bytes(args.size)
	logger.warning(f"Approximate generation time: {round((bytes_size/2097152)/60, 2)} min.")
	generate(data, args.out, bytes_size)
	logger.warning("Generation is finished!")

if __name__ == '__main__':
	main()